from docx import Document #python-docx
import numpy as np
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx_extension import add_page_number, set_cell_border
from text import text_stylization
import pandas as pd
import os
pd.set_option('display.max_columns', None)
pd.set_option('display.max_rows', None)

path = os.getcwd()
new_doc_path = os.path.join(path, 'docs')
if not os.path.exists(new_doc_path):
    os.makedirs(new_doc_path)

file_dir = "files_for_release"
file_name = 'Конструктор.xlsx'

def get_data(file_path:str):
    df = pd.read_excel(file_path, engine='openpyxl', sheet_name='data', dtype={'a': np.str_})
    setting = pd.read_excel(file_path, engine='openpyxl', sheet_name='Settings', index_col=0)
    return df, setting

try:
    df, setting = get_data(file_name)
    dev = 0
except OSError as err:
    df, setting = get_data(file_dir+'/'+file_name)
    dev = 1

documents = df.columns.to_list()

for index in range(len(documents)):
    if documents[index] in ['p', 'RU', 'EN']:
        continue

    df1 = df[(df[documents[index]] == "да")]

    document = Document()

    # update style by default
    styles = document.styles
    styles['Normal'].font.name = 'Times New Roman'
    styles['Normal'].font.size = Pt(12)

    document.sections[0].left_margin = Cm(1)
    document.sections[0].top_margin = Cm(1)
    document.sections[0].right_margin = Cm(1)
    document.sections[0].bottom_margin = Cm(1)

    add_page_number(document.sections[0].footer.paragraphs[0].add_run())
    document.sections[0].footer.paragraphs[0].add_run(' ')
    text_stylization(document.sections[0].footer.paragraphs[0], setting.loc['footer']['value'])

    table = document.add_table(rows=0, cols=3)
    table.style = 'Table Grid'

    for row in df1.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].width = Cm(2)
        row_cells[1].width = Cm(9)
        row_cells[2].width = Cm(9)

        if row[1]['p'] == 'merger_all':
            row_cells[0].merge(row_cells[2])
            text_stylization(row_cells[0].paragraphs[0], row[1]['RU'])
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif str(row[1]['p']) == 'nan':
            row[1]['p'] = ''
            text_stylization(row_cells[0].paragraphs[0], row[1]['p'])
            text_stylization(row_cells[1].paragraphs[0], row[1]['RU'])
            text_stylization(row_cells[2].paragraphs[0], row[1]['EN'])
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            text_stylization(row_cells[0].paragraphs[0], row[1]['p'])
            text_stylization(row_cells[1].paragraphs[0], row[1]['RU'])
            text_stylization(row_cells[2].paragraphs[0], row[1]['EN'])

    #empty row
    document.add_paragraph('')

    #signs
    table_sign = document.add_table(rows=2, cols=2)
    table_sign.style = 'Table Grid'

    table_sign.rows[0].cells[0].merge(table_sign.rows[0].cells[1])
    table_sign.rows[0].cells[0].width = Cm(20)
    text_stylization(table_sign.rows[0].cells[0].paragraphs[0], setting.loc['SignTableName']['value'])
    table_sign.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_sign.rows[0].cells[0].paragraphs[0].runs[0].bold = True

    text_stylization(table_sign.rows[1].cells[0].paragraphs[0], setting.loc['SignBorower']['value'])
    text_stylization(table_sign.rows[1].cells[1].paragraphs[0], setting.loc['SignBusiness']['value'])

    #sign_borrower
    table_sign.rows[1].cells[0].add_paragraph()
    table_sign.rows[1].cells[0].add_paragraph()
    table_sign.rows[1].cells[0].add_paragraph()
    t = table_sign.rows[1].cells[0].add_table(rows=2, cols=1)
    text_stylization(t.rows[0].cells[0].paragraphs[0], setting.loc['SignBorower_sms']['value'])
    t.rows[1].cells[0].paragraphs[0].text = setting.loc['SignBorower_fio']['value']
    set_cell_border(
        t.rows[1].cells[0],
        top={"sz": 1, "val": "single"},
    )

    # sign_too
    t = table_sign.rows[1].cells[1].add_table(rows=2, cols=1)
    if dev == 1:
        t.rows[0].cells[0].paragraphs[0].add_run('').add_picture(
            file_dir+'/'+setting.loc['SignBusiness_pic']['value'],  width=Cm(3), height=Cm(3))
    else:
        t.rows[0].cells[0].paragraphs[0].add_run('').add_picture(
            setting.loc['SignBusiness_pic']['value'], width=Cm(3), height=Cm(3))
    t.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    text_stylization(t.rows[1].cells[0].paragraphs[0], setting.loc['SignBusiness_fio']['value'])
    set_cell_border(
        t.rows[1].cells[0],
        top={"sz": 1, "val": "single"},
    )

    document.save(os.path.join(new_doc_path, documents[index]) + '.docx') # Save document
